# -*- coding: utf-8 -*-
"""Redigerer eksisterende docx in-place: retter navn + deler opp metadatablokken.
Bevarer brukerens avkrysningsbokser (legacy form fields)."""
import os
from docx import Document

DOCX = r"c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell\005 report\Rapport_LOG650_Magnus_Odegard.docx"

NEW_NAME = "Magnus Ødegård Skarsbakk"

doc = Document(DOCX)

# ---------- 1. Forsidefelt "Forfatter(e)": Magnus Ødegård -> fullt navn ----------
for p in doc.paragraphs:
    if p.text.strip() == "Magnus Ødegård":
        for r in p.runs:
            r.text = ""
        (p.runs[0] if p.runs else p.add_run("")).text = NEW_NAME
        break

# ---------- 2. Finn metadatablokken i brodteksten og del den opp ----------
META_LINES = [
    ("Forfatter:", NEW_NAME),
    ("Totalt antall sider inkludert forsiden:", ""),
    ("Sted, innleveringsdato:", "Molde, 2026-05-31"),
    ("Studiepoeng:", "15"),
    ("Veileder:", "Bård Inge Austigard Pettersen"),
]

meta = None
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("Forfatter:") and "Veileder:" in t and "Studiepoeng:" in t:
        meta = p
        break

if meta is None:
    raise SystemExit("Fant ikke metadatablokken!")

style = meta.style
for label, value in META_LINES:
    new_p = meta.insert_paragraph_before()
    new_p.style = style
    run_l = new_p.add_run(label)
    run_l.bold = True
    if value:
        new_p.add_run(" " + value)  # en-space etter etiketten

# fjern den opprinnelige sammenslaatte linjen
meta._p.getparent().remove(meta._p)

doc.save(DOCX)
print("Oppdatert:", DOCX)
