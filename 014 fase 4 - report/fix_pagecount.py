# -*- coding: utf-8 -*-
"""Fyller inn sideantall (39) i 'Totalt antall sider inkludert forsiden' paa
forsiden og i metadatablokken. In-place, bevarer avkrysningsbokser."""
from docx import Document

DOCX = r"c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell\005 report\Rapport_LOG650_Magnus_Odegard.docx"
PAGES = "39"
ENSP = " "  # en-space (samme som malens utfyllingslinje / oevrige metadatalinjer)

doc = Document(DOCX)
front_done = body_done = False
for p in doc.paragraphs:
    t = p.text
    if not t.strip().startswith("Totalt antall sider inkludert forsiden:"):
        continue
    if ENSP in t and not front_done:
        # Forsidefelt: blank utfyllingslinje av en-spaces -> erstatt med tall
        for r in p.runs:
            if ENSP in r.text:
                r.text = ""
        p.add_run(PAGES)
        front_done = True
    elif not body_done:
        # Brodtekst: fet etikett uten verdi -> legg til verdi i samme stil som de andre
        p.add_run(ENSP + PAGES)
        body_done = True

doc.save(DOCX)
print("Forside utfylt:", front_done, "| Brodtekst utfylt:", body_done)
