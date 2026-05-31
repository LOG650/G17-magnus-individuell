# -*- coding: utf-8 -*-
"""Bygger rapport.docx fra rapport_mal.md basert paa LOG650-malen, med auto-TOC."""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell"
TEMPLATE = os.path.join(BASE, "000 templates", "Mal prosjekt LOG650 v2.docx")
MD = os.path.join(BASE, "014 fase 4 - report", "rapport_mal.md")
MD_DIR = os.path.dirname(MD)
OUT_DOCX = os.path.join(BASE, "005 report", "Rapport_LOG650_Magnus_Odegard.docx")

USABLE_WIDTH_IN = 6.0  # tilnaermet tekstbredde A4 med standardmarger

doc = Document(TEMPLATE)

# ---------- 1. Fyll inn forsidefelt ----------
def set_para_text(p, text):
    # behold stil; fjern eksisterende runs, sett ny tekst
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

front = {
    "Tittel (norsk og/eller engelsk)": "Finansiell logistikk og beslutningsstøtte ved hjelp av maskinlæring",
    "Forfatter(e)": "Magnus Ødegård Skarsbakk",
    "Molde, Innleveringsdato": "Molde, 2026-05-31",
}
for p in doc.paragraphs:
    t = p.text.strip()
    if t in front:
        set_para_text(p, front[t])
    elif t.startswith("Studiepoeng:"):
        set_para_text(p, "Studiepoeng:  15")
    elif t.startswith("Veileder:"):
        set_para_text(p, "Veileder:  Bård Inge Austigard Pettersen")

# ---------- 1b. Fjern malens skjelett (manuell TOC + veiledningstekst) ----------
# Behold forsidematerialet; kutt alt fra "Sammendrag"-plassholderen og utover,
# men bevar den avsluttende sectPr (sideoppsett).
body = doc.element.body
found = False
to_remove = []
for child in list(body.iterchildren()):
    tag = child.tag.split('}')[-1]
    if not found and tag == 'p':
        txt = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
        if txt == 'Sammendrag':
            found = True
    if found and tag != 'sectPr':
        to_remove.append(child)
for el in to_remove:
    body.remove(el)

# ---------- 1c. Fjern auto-nummerering fra overskriftsstiler ----------
# Malens Heading-stiler auto-nummererer (1.0, 1.1 ...). Rapporten har egne
# manuelle numre i teksten, saa vi fjerner stilnummereringen for aa unngaa dobbel.
for sname in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
    try:
        st = doc.styles[sname]
    except KeyError:
        continue
    pPr = st.element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)

# ---------- 2. Hjelpere for innhold ----------
def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def add_field(paragraph, instr):
    r = paragraph.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Hoyreklikk og velg \"Oppdater felt\" for innholdsfortegnelse."
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    r._r.append(fb); r._r.append(it); r._r.append(fs)
    r2 = paragraph.add_run(); r2._r.append(t)
    r3 = paragraph.add_run(); r3._r.append(fe)

INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')

def add_runs(paragraph, text):
    text = text.replace(' ', ' ')
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1]); run.font.name = 'Consolas'; run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def add_image(src, width_pct):
    path = os.path.normpath(os.path.join(MD_DIR, src))
    if not os.path.exists(path):
        doc.add_paragraph("[Figur mangler: %s]" % src)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w = Inches(USABLE_WIDTH_IN * (width_pct / 100.0))
    p.add_run().add_picture(path, width=w)

def add_md_table(rows):
    # rows: liste av lister med celletekst; rad 0 = header
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, r in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncol):
            val = r[j] if j < len(r) else ""
            cell = cells[j]
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], val)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

# ---------- 3. Les markdown ----------
with open(MD, encoding='utf-8') as f:
    lines = f.read().split('\n')

# Hopp over toppblokk (tittel/forfatter) og den manuelle Innhold-listen.
# Vi starter behandlingen ved foerste "## " heading; Sammendrag/Abstract beholdes,
# men "## Innhold" + dens liste hoppes over.

# ---------- 4. Sett inn TOC-side ----------
add_page_break()
h = doc.add_paragraph("Innhold")
try:
    h.style = doc.styles['TOC Heading']
except KeyError:
    h.style = doc.styles['Heading 1']
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u ')
add_page_break()

# ---------- 5. Parse hovedinnhold ----------
i = 0
n = len(lines)
skip_section = False  # for "## Innhold"

def parse_figure(block):
    src_m = re.search(r'src="([^"]+)"', block)
    w_m = re.search(r'width="(\d+)%"', block)
    cap_m = re.search(r'<figcaption>\s*<small>(.*?)</small>\s*</figcaption>', block, re.S)
    src = src_m.group(1) if src_m else None
    wpct = int(w_m.group(1)) if w_m else 60
    cap = re.sub(r'\s+', ' ', cap_m.group(1)).strip() if cap_m else ""
    return src, wpct, cap

while i < n:
    line = lines[i]
    stripped = line.strip()

    # Figurblokk
    if stripped.startswith('<figure'):
        block_lines = []
        while i < n and '</figure>' not in lines[i]:
            block_lines.append(lines[i]); i += 1
        if i < n:
            block_lines.append(lines[i]); i += 1
        block = '\n'.join(block_lines)
        src, wpct, cap = parse_figure(block)
        if src:
            add_image(src, wpct)
        if cap:
            add_caption(cap)
        continue

    # Overskrifter
    m = re.match(r'^(#{1,6})\s+(.*)$', line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        plain = re.sub(r'^\d+(\.\d+)*\s*', '', text)  # for sammenligning
        if text == 'Innhold':
            skip_section = True
            i += 1
            continue
        skip_section = False
        # map: ## -> H1, ### -> H2, #### -> H3
        if level == 1:
            i += 1
            continue  # dokumenttittel haandteres av forside
        hlvl = min(level - 1, 3)
        hp = doc.add_paragraph(text)
        hp.style = doc.styles['Heading %d' % hlvl]
        i += 1
        continue

    if skip_section:
        i += 1
        continue

    # Horisontal linje
    if stripped == '---':
        i += 1
        continue

    # Tabell
    if stripped.startswith('|') and i + 1 < n and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i+1]):
        tbl_rows = []
        header = [c.strip() for c in stripped.strip('|').split('|')]
        tbl_rows.append(header)
        i += 2  # hopp header + separator
        while i < n and lines[i].strip().startswith('|'):
            row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            tbl_rows.append(row)
            i += 1
        add_md_table(tbl_rows)
        continue

    # Punktliste
    if re.match(r'^\s*[-*]\s+', line):
        item = re.sub(r'^\s*[-*]\s+', '', line)
        try:
            p = doc.add_paragraph(style='List Paragraph')
        except KeyError:
            p = doc.add_paragraph()
        p.add_run('•  ')
        add_runs(p, item)
        i += 1
        continue

    # Tom linje
    if stripped == '':
        i += 1
        continue

    # Vanlig avsnitt (kan gaa over flere linjer til tom linje)
    para_lines = [stripped]
    i += 1
    while i < n and lines[i].strip() != '' and not lines[i].strip().startswith('|') \
            and not re.match(r'^#{1,6}\s', lines[i]) and not lines[i].strip().startswith('<figure') \
            and lines[i].strip() != '---' and not re.match(r'^\s*[-*]\s+', lines[i]):
        para_lines.append(lines[i].strip())
        i += 1
    text = ' '.join(para_lines)
    p = doc.add_paragraph()
    add_runs(p, text)

doc.save(OUT_DOCX)
print("Lagret:", OUT_DOCX)
