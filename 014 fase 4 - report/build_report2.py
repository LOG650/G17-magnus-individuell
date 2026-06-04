# -*- coding: utf-8 -*-
import os, re, itertools
from xml.sax.saxutils import escape
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT   = r"c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell"
TPL    = os.path.join(ROOT, "000 templates", "Mal prosjekt LOG650 v2.docx")
MD     = os.path.join(ROOT, "014 fase 4 - report", "rapport_mal.md")
MD_DIR = os.path.dirname(MD)
OUT    = os.path.join(ROOT, "005 report", "Rapport LOG650 - Magnus Odegard.docx")

with open(MD, encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document(TPL)
body = doc.element.body

# remove the template's editing restriction (edit="forms"), which otherwise
# locks all normal text so only form fields can be changed in Word
_settings = doc.settings.element
_dp = _settings.find(qn('w:documentProtection'))
if _dp is not None:
    _settings.remove(_dp)

# content width (EMU) for image scaling
sec = doc.sections[0]
CONTENT_W = sec.page_width - sec.left_margin - sec.right_margin

# ---------- inline formatting ----------
TOKEN = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')
def add_runs(p, text):
    for part in TOKEN.split(text):
        if part == "":
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            p.add_run(part)

def set_para_text(el, text, size=None, bold=None):
    """Replace text of an existing template paragraph, keeping its formatting."""
    para = Paragraph(el, None)
    runs = para.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
        if size is not None:
            runs[0].font.size = Pt(size)
        if bold is not None:
            runs[0].bold = bold
    else:
        r = para.add_run(text)
        if size is not None: r.font.size = Pt(size)
        if bold is not None: r.bold = bold

# ---------- 1. fill cover page ----------
children = list(body)
set_para_text(children[14], "Finansiell logistikk og beslutningsstotte ved hjelp av maskinlaering".replace("aelig","").replace("otte","øtte").replace("laering","æring"))
# (build the title cleanly with proper Norwegian chars)
set_para_text(children[14], "Finansiell logistikk og beslutningsstøtte ved hjelp av maskinlæring")
set_para_text(children[16], "Magnus Ødegård")
set_para_text(children[20], "Molde, 2026-05-31")
set_para_text(children[54], "Studiepoeng: 10")
set_para_text(children[55], "Veileder: Bård")

# ---------- 2. remove placeholder body (idx 80..284), keep sectPr (285) ----------
for el in children[80:285]:
    body.remove(el)

# ---------- helpers that append before the trailing sectPr ----------
def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def title_para(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

# ---------- 3. extract Sammendrag & Abstract ----------
def grab_after(header):
    for i, ln in enumerate(lines):
        if ln.strip() == header:
            j = i + 1
            while j < len(lines) and lines[j].strip() in ("", "---"):
                j += 1
            return lines[j].strip()
    return ""

samm = grab_after("## Sammendrag")
abst = grab_after("## Abstract")

page_break()
title_para("Sammendrag")
add_runs(doc.add_paragraph(), samm)
doc.add_paragraph()
title_para("Abstract")
add_runs(doc.add_paragraph(), abst)

# ---------- 4. Table of contents (auto-updating field) ----------
page_break()
title_para("Innhold")

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = ' TOC \\o "1-2" \\h \\z \\u '
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Hoyreklikk her og velg «Oppdater felt» (eller trykk F9) for å generere innholdsfortegnelsen."
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fs)
    run._r.append(placeholder); run._r.append(fe)

add_toc()

# ---------- 5. body parsing ----------
def find_body_start():
    for i, ln in enumerate(lines):
        if ln.strip().startswith("## 1.0"):
            return i
    return 0

def render_heading(s, level):
    s = s.lstrip("#").strip()
    m = re.match(r'^(\d+\.\d+)\s+(.*)$', s)
    text = f"{m.group(1)}\t{m.group(2)}" if m else s
    style = "Heading 1" if level == 1 else "Heading 2"
    doc.add_paragraph(text, style=style)

def render_table(block):
    rows = []
    for ln in block:
        if re.match(r'^\s*\|[\s:|-]+\|\s*$', ln):   # separator row
            continue
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = len(rows[0])
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, cells in enumerate(rows):
        rcells = t.add_row().cells
        for ci in range(ncols):
            txt = cells[ci] if ci < len(cells) else ""
            cell = rcells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if ri == 0:
                r = p.add_run(txt.replace("**", "")); r.bold = True
            else:
                add_runs(p, txt)
    doc.add_paragraph()

def render_figure(block):
    text = "\n".join(block)
    m = re.search(r'<img[^>]*src="([^"]+)"[^>]*?(?:width="(\d+)%")?', text)
    if not m:
        return
    src = m.group(1)
    pct = int(m.group(2)) / 100.0 if m.group(2) else 0.6
    path = os.path.normpath(os.path.join(MD_DIR, src))
    cap = re.search(r'<figcaption>(.*?)</figcaption>', text, re.S)
    caption = ""
    if cap:
        caption = re.sub(r'<[^>]+>', '', cap.group(1)).strip()
    # image
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.isfile(path):
        run = p.add_run()
        run.add_picture(path, width=Emu(int(CONTENT_W * pct)))
    else:
        p.add_run(f"[Bilde mangler: {src}]")
    # caption
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9)
        cp.paragraph_format.space_after = Pt(10)

def render_bullet(s):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Emu(int(0.25 * 914400))
    p.add_run("•\t")
    add_runs(p, s[2:].strip())

i = find_body_start()
n = len(lines)
while i < n:
    raw = lines[i]
    s = raw.strip()
    if s == "" or s == "---":
        i += 1; continue
    if s.startswith("### "):
        render_heading(s, 2); i += 1; continue
    if s.startswith("## "):
        render_heading(s, 1); i += 1; continue
    if s.startswith("|"):
        block = []
        while i < n and lines[i].strip().startswith("|"):
            block.append(lines[i]); i += 1
        render_table(block); continue
    if s.startswith("<figure"):
        block = []
        while i < n and "</figure>" not in lines[i]:
            block.append(lines[i]); i += 1
        if i < n:
            block.append(lines[i]); i += 1   # include closing tag line
        render_figure(block); continue
    if s.startswith("- "):
        render_bullet(s); i += 1; continue
    # default: ordinary paragraph (single line)
    add_runs(doc.add_paragraph(), s)
    i += 1

# ---------- 5b. convert legacy form fields to editable / clickable controls ----------
# Legacy FORMCHECKBOX/FORMTEXT fields only work while the document is protected
# for forms. Since protection is removed (so the body is editable), convert them:
#   checkbox -> modern clickable checkbox content control (toggles ☐/☒ on click)
#   text     -> plain editable text (placeholder underscores if empty)
_id = itertools.count(900100)

def make_checkbox():
    xml = (
        '<w:sdt %s>' % nsdecls('w', 'w14')
        + '<w:sdtPr>'
          '<w:rPr><w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic"/></w:rPr>'
          '<w:id w:val="%d"/>' % next(_id)
        + '<w14:checkbox>'
          '<w14:checked w14:val="0"/>'
          '<w14:checkedState w14:val="2612" w14:font="MS Gothic"/>'
          '<w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>'
          '</w14:checkbox>'
          '</w:sdtPr>'
          '<w:sdtContent>'
          '<w:r><w:rPr><w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic"/></w:rPr>'
          '<w:t>☐</w:t></w:r>'
          '</w:sdtContent>'
          '</w:sdt>'
    )
    return parse_xml(xml)

def make_text_run(text):
    text = text if text.strip() else "_____"
    return parse_xml('<w:r %s><w:t xml:space="preserve">%s</w:t></w:r>' % (nsdecls('w'), escape(text)))

def make_dropdown(entries):
    if not entries:
        return make_text_run("")
    # pre-select the LOG650 entry if present, otherwise show the first (placeholder)
    shown = next((e for e in entries if "LOG650" in e), entries[0])
    items = "".join(
        '<w:listItem w:displayText="%s" w:value="%s"/>' % (escape(e, {'"': '&quot;'}), escape(e, {'"': '&quot;'}))
        for e in entries
    )
    xml = (
        '<w:sdt %s>' % nsdecls('w')
        + '<w:sdtPr><w:id w:val="%d"/>' % next(_id)
        + '<w:dropDownList>%s</w:dropDownList>' % items
        + '</w:sdtPr>'
        + '<w:sdtContent><w:r><w:t xml:space="preserve">%s</w:t></w:r></w:sdtContent>' % escape(shown)
        + '</w:sdt>'
    )
    return parse_xml(xml)

def convert_fields(paragraph_el):
    children = list(paragraph_el)
    n = len(children)
    groups = []
    i = 0
    while i < n:
        ch = children[i]
        if ch.tag == qn('w:r'):
            fc = ch.find(qn('w:fldChar'))
            if fc is not None and fc.get(qn('w:fldCharType')) == 'begin':
                ftype = None; result = []; seen_sep = False; j = i
                entries = [le.get(qn('w:val')) for le in ch.findall('.//' + qn('w:listEntry'))]
                while j < n:
                    cj = children[j]
                    if cj.tag == qn('w:r'):
                        instr = cj.find(qn('w:instrText'))
                        if instr is not None and instr.text:
                            if 'FORMCHECKBOX' in instr.text: ftype = 'checkbox'
                            elif 'FORMDROPDOWN' in instr.text: ftype = 'dropdown'
                            elif 'FORMTEXT' in instr.text: ftype = 'text'
                        fcj = cj.find(qn('w:fldChar'))
                        if fcj is not None:
                            ct = fcj.get(qn('w:fldCharType'))
                            if ct == 'separate': seen_sep = True
                            elif ct == 'end':
                                if ftype in ('checkbox', 'text', 'dropdown'):   # leave TOC & other fields intact
                                    groups.append((i, j, ftype, ''.join(result), entries))
                                break
                        elif seen_sep:
                            t = cj.find(qn('w:t'))
                            if t is not None and t.text: result.append(t.text)
                    j += 1
                i = j + 1; continue
        i += 1
    for start, end, ftype, result, entries in reversed(groups):
        if ftype == 'checkbox':
            repl = make_checkbox()
        elif ftype == 'dropdown':
            repl = make_dropdown(entries)
        else:
            repl = make_text_run(result)
        children[start].addprevious(repl)
        for k in range(start, end + 1):
            paragraph_el.remove(children[k])

for _p in body.iter(qn('w:p')):
    if _p.find('.//' + qn('w:fldChar')) is not None:
        convert_fields(_p)

# ---------- 6. tell Word to update fields (TOC) on open ----------
settings = doc.settings.element
uf = settings.find(qn('w:updateFields'))
if uf is None:
    uf = OxmlElement('w:updateFields')
    settings.append(uf)
uf.set(qn('w:val'), 'true')

# page break before chapter 1 so TOC sits on its own page
# (insert a page break paragraph right before the first Heading 1)
for p in doc.paragraphs:
    if p.style.name == "Heading 1":
        brk = OxmlElement('w:p')
        run = OxmlElement('w:r'); br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
        run.append(br); brk.append(run)
        p._element.addprevious(brk)
        break

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("SAVED:", OUT)
print("Sammendrag chars:", len(samm), "| Abstract chars:", len(abst))
