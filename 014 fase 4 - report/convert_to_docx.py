#!/usr/bin/env python3
"""Convert rapport_mal.md to Word document using LOG650 template."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = r'c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell\000 templates\Mal prosjekt LOG650 v2.docx'
MD_FILE  = r'c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell\014 fase 4 - report\rapport_mal.md'
OUT_FILE = r'c:\Users\mrmag\Desktop\SKOLE\LOG650\G17-magnus-individuell\014 fase 4 - report\rapport_magnus_odegard.docx'
MD_DIR   = os.path.dirname(MD_FILE)

FONT = 'Arial'

# Norwegian Word style IDs (localized template)
S_NORMAL  = 'Normal'
S_H1      = 'Overskrift1'
S_H2      = 'Overskrift2'
S_H3      = 'Overskrift3'
S_H4      = 'Overskrift4'
S_TOC1    = 'INNH1'
S_TOC2    = 'INNH2'
S_LIST    = 'Listeavsnitt'

# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _new_p(style_id=None):
    p = OxmlElement('w:p')
    if style_id:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    return p


def _new_r(text, bold=False, italic=False, font_size_pt=12):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)
    sz_val = str(int(font_size_pt * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    r.append(rPr)
    t = OxmlElement('w:t')
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def _add_formatted_runs(p_elem, text, font_size_pt=12):
    """Add runs with **bold** and *italic* markdown support."""
    # Handle escaped text first
    parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            p_elem.append(_new_r(part[2:-2], bold=True, font_size_pt=font_size_pt))
        elif part.startswith('*') and part.endswith('*'):
            p_elem.append(_new_r(part[1:-1], italic=True, font_size_pt=font_size_pt))
        else:
            p_elem.append(_new_r(part, font_size_pt=font_size_pt))


def _set_center(p_elem):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)


def delete_paragraph(para):
    elem = para._element
    elem.getparent().remove(elem)


# ─────────────────────────────────────────────────────────────────────────────
# Cover-page helpers
# ─────────────────────────────────────────────────────────────────────────────

def replace_para_text(para, text, bold=False, font_size_pt=12):
    para.clear()
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(font_size_pt)
    run.bold = bold


# ─────────────────────────────────────────────────────────────────────────────
# Markdown parser
# ─────────────────────────────────────────────────────────────────────────────

def parse_md(md_text):
    lines = md_text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if re.match(r'^#### ', line):
            blocks.append({'type': 'h4', 'text': line[5:].strip()})
        elif re.match(r'^### ', line):
            blocks.append({'type': 'h3', 'text': line[4:].strip()})
        elif re.match(r'^## ', line):
            blocks.append({'type': 'h2', 'text': line[3:].strip()})
        elif re.match(r'^# ', line):
            blocks.append({'type': 'h1_doc', 'text': line[2:].strip()})
        elif line.strip() == '---':
            blocks.append({'type': 'hr'})
        elif line.strip().startswith('<figure'):
            fig_lines = [line]
            i += 1
            while i < len(lines) and '</figure>' not in lines[i]:
                fig_lines.append(lines[i])
                i += 1
            if i < len(lines):
                fig_lines.append(lines[i])
            all_fig = '\n'.join(fig_lines)
            src_m  = re.search(r'src=["\']([^"\']+)["\']', all_fig)
            cap_m  = re.search(r'<figcaption><small>(.*?)</small></figcaption>', all_fig, re.DOTALL)
            blocks.append({
                'type': 'figure',
                'src': src_m.group(1) if src_m else None,
                'caption': cap_m.group(1).strip() if cap_m else ''
            })
        elif line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[-: |]+\|', lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'table', 'lines': table_lines})
            continue
        elif re.match(r'^( {0,6})[-*] ', line):
            indent = len(re.match(r'^( *)', line).group(1))
            level = 1 + indent // 2
            text = re.sub(r'^ *[-*] ', '', line)
            blocks.append({'type': 'bullet', 'text': text.strip(), 'level': level})
        elif re.match(r'^\d+\.', line.strip()) and not line.strip().startswith('0.'):
            blocks.append({'type': 'toc_item', 'text': line.strip()})
        elif line.strip() == '':
            blocks.append({'type': 'empty'})
        elif line.startswith('**') and ':**' in line:
            blocks.append({'type': 'cover_kv', 'text': line})
        else:
            blocks.append({'type': 'para', 'text': line})

        i += 1

    return blocks


def get_section(blocks, heading_text):
    """Get blocks between an h2 heading and the next h2/hr."""
    in_section = False
    result = []
    for b in blocks:
        if b['type'] == 'h2' and b['text'] == heading_text:
            in_section = True
            continue
        if in_section:
            if b['type'] in ('h2', 'hr'):
                break
            result.append(b)
    return result


def get_text_blocks(blocks):
    return [b['text'].strip() for b in blocks if b['type'] == 'para' and b['text'].strip()]


# ─────────────────────────────────────────────────────────────────────────────
# DocWriter – appends content before final body sectPr
# ─────────────────────────────────────────────────────────────────────────────

class DocWriter:
    def __init__(self, doc):
        self.doc = doc
        self.body = doc.element.body

    def _append(self, elem):
        body_sectPr = self.body.find(qn('w:sectPr'))
        if body_sectPr is not None:
            body_sectPr.addprevious(elem)
        else:
            self.body.append(elem)

    def add_para(self, text, style=S_NORMAL, bold=False, italic=False, center=False):
        p = _new_p(style)
        if center:
            _set_center(p)
        _add_formatted_runs(p, text)
        self._append(p)

    def add_bold_label(self, text):
        """Add a paragraph like '**Tabell 4.1 – Nøkkeltall**' as centered bold."""
        p = _new_p(S_NORMAL)
        p.append(_new_r(text, bold=True))
        self._append(p)

    def add_heading(self, text, level=1):
        style_map = {1: S_H1, 2: S_H2, 3: S_H3, 4: S_H4}
        p = _new_p(style_map.get(level, S_H1))
        p.append(_new_r(text))
        self._append(p)

    def add_empty(self):
        p = _new_p(S_NORMAL)
        self._append(p)

    def add_bullet(self, text, level=1):
        p = _new_p(S_LIST)
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        # Indentation
        ind = OxmlElement('w:ind')
        left_twips = str(720 * level)
        ind.set(qn('w:left'), left_twips)
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        # Bullet character via numPr if numbering is defined, else use indent only
        _add_formatted_runs(p, '– ' + text)
        self._append(p)

    def add_table(self, rows):
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        # Equal column widths in percentage (5000 = 100%)
        col_w = max(500, 5000 // num_cols)

        tbl = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblPr.append(tblStyle)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)
        # Borders
        tblBorders = OxmlElement('w:tblBorders')
        for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bname}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(col_w))
            tblGrid.append(gc)
        tbl.append(tblGrid)

        for row_i, row in enumerate(rows):
            tr = OxmlElement('w:tr')
            is_header = (row_i == 0)
            for col_i in range(num_cols):
                cell_text = row[col_i] if col_i < len(row) else ''
                tc = OxmlElement('w:tc')
                tcPr = OxmlElement('w:tcPr')
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(col_w))
                tcW.set(qn('w:type'), 'pct')
                tcPr.append(tcW)
                if is_header:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'D9D9D9')
                    tcPr.append(shd)
                tc.append(tcPr)
                cp = _new_p(S_NORMAL)
                if is_header:
                    cp.append(_new_r(cell_text, bold=True, font_size_pt=11))
                else:
                    _add_formatted_runs(cp, cell_text, font_size_pt=11)
                tc.append(cp)
                tr.append(tc)
            tbl.append(tr)

        self._append(tbl)
        self.add_empty()

    def add_image(self, src, caption=''):
        if not src:
            return
        img_path = os.path.normpath(os.path.join(MD_DIR, src.replace('/', os.sep)))
        if not os.path.exists(img_path):
            self.add_para(f'[Figur: {os.path.basename(img_path)}]', italic=True)
            if caption:
                self.add_para(caption, italic=True)
            return

        self.add_empty()

        # Add image via python-docx (it appends to body, then we move it)
        temp_para = self.doc.add_paragraph()
        temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = temp_para.add_run()
        try:
            run.add_picture(img_path, width=Inches(4.5))
        except Exception as e:
            print(f'  Warning image {os.path.basename(img_path)}: {e}')
            temp_para._element.getparent().remove(temp_para._element)
            self.add_para(f'[Figur: {os.path.basename(img_path)}]', italic=True)
            if caption:
                self.add_para(caption, italic=True)
            return

        # doc.add_paragraph() appends at end; move it to the correct position
        tp_elem = temp_para._element
        tp_elem.getparent().remove(tp_elem)
        _set_center(tp_elem)
        self._append(tp_elem)

        if caption:
            cap_p = _new_p(S_NORMAL)
            _set_center(cap_p)
            cap_p.append(_new_r(caption, italic=True, font_size_pt=10))
            self._append(cap_p)

        self.add_empty()

    def add_page_break(self):
        p = _new_p(S_NORMAL)
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        self._append(p)


# ─────────────────────────────────────────────────────────────────────────────
# Markdown table parser
# ─────────────────────────────────────────────────────────────────────────────

def parse_table_lines(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\|[-: |]+\|$', stripped):
            continue  # skip separator
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        if any(cells):
            rows.append(cells)
    return rows


# ─────────────────────────────────────────────────────────────────────────────
# Block renderer
# ─────────────────────────────────────────────────────────────────────────────

def render_blocks(writer, blocks):
    prev_empty = False
    for blk in blocks:
        btype = blk['type']

        if btype == 'h2':
            text = blk['text']
            m = re.match(r'^(\d+\.(?:\d+)?)\s+(.+)$', text)
            heading_text = f'{m.group(1)}\t{m.group(2)}' if m else text
            writer.add_heading(heading_text, level=1)
            prev_empty = False

        elif btype == 'h3':
            text = blk['text']
            m = re.match(r'^(\d+\.\d+)\s+(.+)$', text)
            heading_text = f'{m.group(1)}\t{m.group(2)}' if m else text
            writer.add_heading(heading_text, level=2)
            prev_empty = False

        elif btype == 'h4':
            writer.add_heading(blk['text'], level=3)
            prev_empty = False

        elif btype == 'para':
            text = blk['text'].strip()
            if not text:
                continue
            # Bold standalone labels (table/figure titles)
            if re.match(r'^\*\*(?:Tabell|Figur|Table)\s', text) and text.endswith('**'):
                writer.add_bold_label(text[2:-2])
            else:
                writer.add_para(text)
            prev_empty = False

        elif btype == 'bullet':
            writer.add_bullet(blk['text'], blk['level'])
            prev_empty = False

        elif btype == 'figure':
            writer.add_image(blk['src'], blk['caption'])
            prev_empty = False

        elif btype == 'table':
            rows = parse_table_lines(blk['lines'])
            writer.add_table(rows)
            prev_empty = False

        elif btype == 'empty':
            if not prev_empty:
                writer.add_empty()
            prev_empty = True
            continue

        elif btype == 'hr':
            pass  # Section dividers handled by headings

        if btype != 'empty':
            prev_empty = False


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_text = f.read()

    blocks = parse_md(md_text)

    sammendrag_texts = get_text_blocks(get_section(blocks, 'Sammendrag'))
    abstract_texts   = get_text_blocks(get_section(blocks, 'Abstract'))

    body_start = next(
        (i for i, b in enumerate(blocks)
         if b['type'] == 'h2' and re.match(r'^1\.0\s+', b['text'])), None)
    body_blocks = blocks[body_start:] if body_start is not None else []

    # ── Open template ────────────────────────────────────────────────────────
    doc = Document(TEMPLATE)

    # ── Update cover page ────────────────────────────────────────────────────
    replace_para_text(doc.paragraphs[13],
        'Finansiell logistikk og beslutningsstøtte ved hjelp av maskinlæring',
        bold=True)
    replace_para_text(doc.paragraphs[15], 'Magnus Ødegård')
    replace_para_text(doc.paragraphs[17], 'Totalt antall sider inkludert forsiden: ')
    replace_para_text(doc.paragraphs[19], 'Molde, 2026-05-31')
    replace_para_text(doc.paragraphs[52], 'Studiepoeng: 10', bold=True)
    replace_para_text(doc.paragraphs[53], 'Veileder: Bård', bold=True)

    # ── Insert Sammendrag text ────────────────────────────────────────────────
    sammendrag_para = next(
        (p for p in doc.paragraphs if p.text.strip() == 'Sammendrag'), None)
    if sammendrag_para and sammendrag_texts:
        for text in reversed(sammendrag_texts):
            p = _new_p(S_NORMAL)
            _add_formatted_runs(p, text)
            sammendrag_para._element.addnext(p)

    # ── Insert Abstract text ──────────────────────────────────────────────────
    abstract_para = next(
        (p for p in doc.paragraphs if p.text.strip() == 'Abstract'), None)
    if abstract_para and abstract_texts:
        for text in reversed(abstract_texts):
            p = _new_p(S_NORMAL)
            _add_formatted_runs(p, text)
            abstract_para._element.addnext(p)

    # ── Update TOC ───────────────────────────────────────────────────────────
    # Delete ALL old toc 1/toc 2 paragraphs (template placeholders)
    for p in list(doc.paragraphs):
        if p.style.name in ('toc 1', 'toc 2') and p.text.strip() != 'Innhold':
            delete_paragraph(p)

    innhold_para = next(
        (p for p in doc.paragraphs if p.style.name == 'toc 1'), None)

    toc_entries = [
        (1, '1.0\tInnledning'),
        (2, '1.1\tForskningsmål og problemstilling'),
        (2, '1.2\tDelproblemer'),
        (2, '1.3\tAvgrensinger'),
        (2, '1.4\tAntagelser'),
        (1, '2.0\tLitteratur'),
        (2, '2.1\tLitteratursøk og utvalg'),
        (2, '2.2\tAppel et al. (2019)'),
        (2, '2.3\tSchoonbee et al. (2022)'),
        (2, '2.4\tSammenstilling, forskningsgap og relevans'),
        (1, '3.0\tTeori'),
        (2, '3.1\tMaskinlæring og klassifisering'),
        (2, '3.2\tKandidatalgoritmer'),
        (2, '3.3\tEvalueringsmetrikker'),
        (2, '3.4\tFeature engineering'),
        (2, '3.5\tKonseptdrift'),
        (2, '3.6\tBeslutningsstøttesystem (DSS) og risikoscore'),
        (1, '4.0\tCasebeskrivelse'),
        (1, '5.0\tMetode og data'),
        (2, '5.1\tMetode'),
        (2, '5.2\tData'),
        (2, '5.3\tValiditet, reliabilitet og etiske hensyn'),
        (1, '6.0\tModellering'),
        (2, '6.1\tFeature engineering'),
        (2, '6.2\tBaseline-modeller'),
        (2, '6.3\tHyperparameterjustering'),
        (1, '7.0\tAnalyse'),
        (2, '7.1\tEksplorativ dataanalyse (EDA)'),
        (2, '7.2\tLeverandørprofil og risikoscore'),
        (1, '8.0\tResultat'),
        (2, '8.1\tModellytelse'),
        (2, '8.2\tKonfusjonsmatrise og feature importance'),
        (2, '8.3\tRisikoklassifisering av fakturaer'),
        (1, '9.0\tDiskusjon'),
        (2, '9.1\tModellytelse mot benchmark'),
        (2, '9.2\tBeslutningsstøtteverdi utover dagens praksis'),
        (2, '9.3\tBegrensninger og veien videre'),
        (2, '9.4\tTeoretiske og policy-implikasjoner'),
        (1, '10.0\tKonklusjon'),
        (1, '11.0\tBibliografi'),
        (1, '12.0\tVedlegg'),
    ]

    if innhold_para:
        for level, text in reversed(toc_entries):
            style_id = S_TOC1 if level == 1 else S_TOC2
            p = _new_p(style_id)
            p.append(_new_r(text, bold=(level == 1)))
            innhold_para._element.addnext(p)

    # ── Remove template chapter placeholder content ────────────────────────────
    body = doc.element.body
    body_children = list(body)

    # Find the section-break paragraph that follows the TOC entries
    # (this is the paragraph with sectPr that ends the TOC section)
    toc_sectpr_elem = None
    for child in body_children:
        if child.tag == qn('w:p'):
            if child.find('.//' + qn('w:sectPr')) is not None:
                # Check if this is after any toc entry
                child_idx = list(body).index(child)
                # We'll find the one right after the toc block
                toc_sectpr_elem = child

    # Find all Heading 1 paragraphs (template placeholders) and delete everything
    # from first template heading onwards
    first_template_h1 = None
    for child in list(body):
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == S_H1:
                    first_template_h1 = child
                    break

    if first_template_h1 is not None:
        current_body = list(body)
        h1_idx = current_body.index(first_template_h1)
        for child in current_body[h1_idx:]:
            if child.tag == qn('w:p'):
                # Keep sectPr paragraphs (they define section layout)
                if child.find('.//' + qn('w:sectPr')) is not None:
                    # Clear text but keep the element
                    for t_elem in child.findall('.//' + qn('w:t')):
                        t_elem.text = ''
                    for r_elem in child.findall(qn('w:r')):
                        # Remove runs that don't have section properties
                        if r_elem.find('.//' + qn('w:sectPr')) is None:
                            pPr_elem = child.find(qn('w:pPr'))
                            if r_elem.getparent() == child:
                                child.remove(r_elem)
                    continue
                body.remove(child)
            elif child.tag == qn('w:tbl'):
                body.remove(child)

    # ── Write chapter content ─────────────────────────────────────────────────
    writer = DocWriter(doc)
    render_blocks(writer, body_blocks)

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(OUT_FILE)
    print(f'Saved: {OUT_FILE}')
    print(f'  Paragraphs: {len(doc.paragraphs)}')
    print(f'  Tables: {len(doc.tables)}')

    # Quick summary
    h1s = [p.text[:60] for p in doc.paragraphs if p.style.name == 'Heading 1']
    print(f'  Heading 1 count: {len(h1s)}')
    for h in h1s:
        print(f'    {h}')


if __name__ == '__main__':
    main()
